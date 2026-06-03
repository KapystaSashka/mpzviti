import datetime
import re
import jwt
import io
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

from app.database import db, User, Event, AuditLog, ROLES, EVENT_CATEGORIES
from app.report_structure import REPORT_STRINGS

SECRET_KEY = "viti_secret_mpz_key_system_2026_secured_key"


def has_permission(role: str, permission: str) -> bool:
    perms = ROLES.get(role, [])
    return "*" in perms or permission in perms


class AuthService:
    @staticmethod
    def login(username: str, password: str):
        user = User.query.filter_by(username=username, password=password).first()
        if user:
            payload = {
                "sub":  user.username,
                "role": user.role,
                "name": user.name,
                "unit": user.unit,
                "exp":  datetime.datetime.utcnow() + datetime.timedelta(hours=8),
            }
            return jwt.encode(payload, SECRET_KEY, algorithm="HS256")
        return None

    @staticmethod
    def verify_token(token: str):
        try:
            return jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
            return None


class EventService:
    @staticmethod
    def get_all_events(unit_id=None):
        q = Event.query
        if unit_id is not None:
            q = q.filter((Event.unit_id == int(unit_id)) | (Event.unit_id == None))
        return [e.to_dict() for e in q.order_by(Event.date.desc()).all()]

    @staticmethod
    def add_event(data: dict, username: str = "system"):
        ev = Event(
            title=data.get("title", ""), date=data.get("date", ""),
            location=data.get("location", ""), responsible=data.get("responsible", ""),
            status=data.get("status", "Заплановано"), category=data.get("category", "cultural"),
            unit_id=data.get("unit_id"), attendance=int(data.get("attendance", 0)),
            capacity=int(data.get("capacity", 0)),
            satisfaction_score=float(data.get("satisfaction_score", 0)),
            notes=data.get("notes", ""),
        )
        db.session.add(ev)
        db.session.flush()
        db.session.add(AuditLog(username=username, action="create", entity="event",
                                entity_id=ev.id, detail=f"Створено захід: «{ev.title}»"))
        db.session.commit()
        return ev.to_dict()

    @staticmethod
    def update_event(event_id: int, data: dict, username: str = "system"):
        ev = Event.query.get(event_id)
        if not ev:
            return None
        changed = []
        for field in ["title","date","location","responsible","status","category",
                      "unit_id","attendance","capacity","satisfaction_score","notes"]:
            if field in data:
                old = getattr(ev, field)
                if str(old) != str(data[field]):
                    changed.append(f"{field}: «{old}» → «{data[field]}»")
                setattr(ev, field, data[field])
        if changed:
            db.session.add(AuditLog(username=username, action="update", entity="event",
                                    entity_id=ev.id, detail="; ".join(changed)))
        db.session.commit()
        return ev.to_dict()

    @staticmethod
    def delete_event(event_id: int, username: str = "system"):
        ev = Event.query.get(event_id)
        if not ev:
            return None
        db.session.add(AuditLog(username=username, action="delete", entity="event",
                                entity_id=ev.id, detail=f"Видалено захід: «{ev.title}»"))
        db.session.delete(ev)
        db.session.commit()
        return {"id": event_id}


class AuditService:
    @staticmethod
    def get_log(limit: int = 50):
        logs = AuditLog.query.order_by(AuditLog.created_at.desc()).limit(limit).all()
        return [l.to_dict() for l in logs]


class AnalyticsService:
    @staticmethod
    def get_summary():
        events = Event.query.all()
        completed = [e for e in events if e.status == "Завершено"]
        avg_sat = avg_att = 0.0
        if completed:
            scored = [e for e in completed if e.satisfaction_score > 0]
            avg_sat = round(sum(e.satisfaction_score for e in scored)/len(scored), 2) if scored else 0
            filled = [e for e in completed if e.capacity > 0]
            avg_att = round(sum(e.attendance/e.capacity*100 for e in filled)/len(filled), 1) if filled else 0
        return {"total_events": len(events), "completed_events": len(completed),
                "planned_events": len([e for e in events if e.status=="Заплановано"]),
                "avg_satisfaction": avg_sat, "avg_attendance_pct": avg_att}

    @staticmethod
    def get_category_stats():
        events = Event.query.all()
        stats = defaultdict(lambda: {"count":0,"total_satisfaction":0.0,"total_attendance":0,"total_capacity":0})
        for ev in events:
            cat = ev.category or "other"
            stats[cat]["count"] += 1
            stats[cat]["total_satisfaction"] += ev.satisfaction_score or 0
            stats[cat]["total_attendance"]   += ev.attendance or 0
            stats[cat]["total_capacity"]     += ev.capacity or 0
        result = []
        for cat, data in stats.items():
            c = data["count"]
            result.append({"category": cat, "category_label": EVENT_CATEGORIES.get(cat, cat),
                           "event_count": c,
                           "avg_satisfaction": round(data["total_satisfaction"]/c, 2) if c else 0,
                           "avg_attendance_pct": round(data["total_attendance"]/data["total_capacity"]*100, 1) if data["total_capacity"] > 0 else 0})
        return sorted(result, key=lambda x: x["avg_satisfaction"], reverse=True)

    @staticmethod
    def get_activity_by_month():
        month_count = defaultdict(int)
        for ev in Event.query.all():
            try: month_count[ev.date[:7]] += 1
            except: pass
        return [{"month": k, "count": v} for k, v in sorted(month_count.items())]

    @staticmethod
    def get_unit_activity():
        events = Event.query.all()
        units = [{"id":1,"name":"1-й факультет"},{"id":2,"name":"2-й факультет"},
                 {"id":3,"name":"3-й факультет"},{"id":4,"name":"Відділ МПЗ"},{"id":5,"name":"Керівництво"}]
        unit_stats = defaultdict(lambda: {"event_count":0,"total_attendance":0})
        for ev in events:
            uid = ev.unit_id
            targets = [1,2,3,4,5] if uid is None else [uid]
            for t in targets:
                unit_stats[t]["event_count"] += 1
                unit_stats[t]["total_attendance"] += ev.attendance or 0
        result = [{"unit_id":u["id"],"unit_name":u["name"],
                   "event_count":unit_stats[u["id"]]["event_count"],
                   "total_attendance":unit_stats[u["id"]]["total_attendance"]} for u in units]
        return sorted(result, key=lambda x: x["event_count"], reverse=True)


class AIService:
    @staticmethod
    def get_rule_based_recommendations():
        recs = []
        cat_stats  = AnalyticsService.get_category_stats()
        unit_stats = AnalyticsService.get_unit_activity()
        for cat in cat_stats:
            if 0 < cat["avg_satisfaction"] < 4.0:
                recs.append({"type":"warning","message":f"Заходи категорії «{cat['category_label']}» мають низький рівень задоволеності ({cat['avg_satisfaction']}/5.0). Рекомендується переглянути формат."})
        if unit_stats:
            mx = unit_stats[0]["event_count"]
            for u in unit_stats:
                if mx > 0 and u["event_count"] < mx * 0.5:
                    recs.append({"type":"suggestion","message":f"Підрозділ «{u['unit_name']}» бере участь у значно меншій кількості заходів ({u['event_count']} vs {mx}). Рекомендується провести додатковий захід."})
        if cat_stats and cat_stats[0]["avg_satisfaction"] >= 4.5:
            best = cat_stats[0]
            recs.append({"type":"positive","message":f"Заходи категорії «{best['category_label']}» показують найвищу ефективність ({best['avg_satisfaction']}/5.0). Рекомендується збільшити частоту."})
        today = datetime.date.today()
        upcoming = [e for e in Event.query.all() if e.status=="Заплановано"
                    and 0 <= (datetime.date.fromisoformat(e.date)-today).days <= 7]
        if not upcoming:
            recs.append({"type":"info","message":"На найближчі 7 днів немає запланованих заходів. Рекомендується внести до плану хоча б один захід."})
        return recs

    @staticmethod
    def get_ai_recommendations(use_claude_api: bool = False):
        rule_recs = AIService.get_rule_based_recommendations()
        summary   = AnalyticsService.get_summary()
        if not use_claude_api:
            return {"mode":"rule_based","recommendations":rule_recs,"summary":summary}
        cat_stats  = AnalyticsService.get_category_stats()
        unit_stats = AnalyticsService.get_unit_activity()
        prompt = f"""Ти — аналітик МПЗ Військового інституту. Статистика:
- Заходів: {summary['total_events']}, завершено: {summary['completed_events']}, заплановано: {summary['planned_events']}
- Задоволеність: {summary['avg_satisfaction']}/5.0, явка: {summary['avg_attendance_pct']}%
Категорії: {'; '.join(f"{c['category_label']} {c['avg_satisfaction']}/5 явка {c['avg_attendance_pct']}%" for c in cat_stats)}
Підрозділи: {'; '.join(f"{u['unit_name']} {u['event_count']} заходів" for u in unit_stats)}
Надай 3-5 рекомендацій JSON: [{{"type":"suggestion|warning|positive","message":"..."}}]. ТІЛЬКИ JSON."""
        return {"mode":"claude_api","context_prompt":prompt,"rule_recommendations":rule_recs,"summary":summary}



# ---------------------------------------------------------------------------
# UserService — управління профілями (тільки admin)
# ---------------------------------------------------------------------------

class UserService:

    # Валідація
    LETTERS_ONLY = re.compile(r"^[А-ЯІЇЄҐа-яіїєґA-Za-z'\-\s]+$")
    USERNAME_RE  = re.compile(r"^[a-zA-Z][a-zA-Z0-9_]{2,31}$")
    PASSWORD_RE  = re.compile(r"^(?=.*[A-Za-z])(?=.*\d).{6,}$")

    VALID_ROLES  = {"admin", "psychologist", "commander", "staff"}
    VALID_RANKS  = [
        "", "рядовий", "молодший сержант", "сержант", "старший сержант",
        "старшина", "прапорщик", "старший прапорщик",
        "молодший лейтенант", "лейтенант", "старший лейтенант", "капітан",
        "майор", "підполковник", "полковник",
        "генерал-майор", "генерал-лейтенант", "генерал-полковник", "генерал армії",
    ]

    @staticmethod
    def validate_user_data(data: dict, is_update: bool = False) -> list:
        """Повертає список помилок або порожній список якщо все ок."""
        errors = []

        if not is_update:
            username = data.get("username", "").strip()
            if not username:
                errors.append("Логін є обов'язковим")
            elif not UserService.USERNAME_RE.match(username):
                errors.append("Логін: лише латиниця, цифри та '_', починається з букви, 3–32 символи")
            elif User.query.filter_by(username=username).first():
                errors.append(f"Логін «{username}» вже зайнятий")

            password = data.get("password", "")
            if not password:
                errors.append("Пароль є обов'язковим")
            elif not UserService.PASSWORD_RE.match(password):
                errors.append("Пароль: мінімум 6 символів, має містити букви та цифри")

        first_name = data.get("first_name", "").strip()
        last_name  = data.get("last_name", "").strip()

        if not is_update or first_name:
            if not first_name:
                errors.append("Ім'я є обов'язковим")
            elif not UserService.LETTERS_ONLY.match(first_name):
                errors.append("Ім'я: лише букви та дефіс")

        if not is_update or last_name:
            if not last_name:
                errors.append("Прізвище є обов'язковим")
            elif not UserService.LETTERS_ONLY.match(last_name):
                errors.append("Прізвище: лише букви та дефіс")

        role = data.get("role", "")
        if role and role not in UserService.VALID_ROLES:
            errors.append(f"Роль «{role}» не існує")

        gender = data.get("gender", "")
        if gender and gender not in ("M", "F"):
            errors.append("Стать: вкажіть M або F")

        return errors

    @staticmethod
    def get_all_users():
        return [u.to_dict() for u in User.query.order_by(User.last_name).all()]

    @staticmethod
    def create_user(data: dict, created_by: str = "system"):
        errors = UserService.validate_user_data(data)
        if errors:
            return None, errors

        user = User(
            username   = data["username"].strip(),
            password   = data["password"],
            role       = data.get("role", "staff"),
            first_name = data.get("first_name", "").strip(),
            last_name  = data.get("last_name", "").strip(),
            rank       = data.get("rank", ""),
            position   = data.get("position", "").strip(),
            gender     = data.get("gender", "M"),
            unit       = data.get("unit", "").strip(),
            is_active  = True,
        )
        db.session.add(user)
        db.session.flush()
        db.session.add(AuditLog(
            username=created_by, action="create", entity="user",
            entity_id=user.id,
            detail=f"Створено користувача: «{user.username}» ({user.name}), роль: {user.role}",
        ))
        db.session.commit()
        return user.to_dict(), []

    @staticmethod
    def update_user(user_id: int, data: dict, updated_by: str = "system"):
        user = User.query.get(user_id)
        if not user:
            return None, ["Користувача не знайдено"]

        errors = UserService.validate_user_data(data, is_update=True)
        if errors:
            return None, errors

        fields = ["first_name","last_name","rank","position","gender","unit","role","is_active"]
        changed = []
        for f in fields:
            if f in data:
                old = getattr(user, f)
                if str(old) != str(data[f]):
                    changed.append(f"{f}: «{old}» → «{data[f]}»")
                setattr(user, f, data[f])

        if "password" in data and data["password"]:
            if not UserService.PASSWORD_RE.match(data["password"]):
                return None, ["Пароль: мінімум 6 символів, має містити букви та цифри"]
            user.password = data["password"]
            changed.append("пароль змінено")

        if changed:
            db.session.add(AuditLog(
                username=updated_by, action="update", entity="user",
                entity_id=user.id, detail="; ".join(changed),
            ))
        db.session.commit()
        return user.to_dict(), []

    @staticmethod
    def delete_user(user_id: int, deleted_by: str = "system"):
        user = User.query.get(user_id)
        if not user:
            return False
        db.session.add(AuditLog(
            username=deleted_by, action="delete", entity="user",
            entity_id=user.id,
            detail=f"Видалено користувача: «{user.username}» ({user.name})",
        ))
        db.session.delete(user)
        db.session.commit()
        return True


# ---------------------------------------------------------------------------
# QRService — генерація QR-кодів для заходів
# ---------------------------------------------------------------------------

class QRService:
    @staticmethod
    def generate_qr(event_id: int, base_url: str = ""):
        """Генерує QR-код для реєстрації відвідування заходу."""
        import qrcode
        import base64

        attend_url = f"{base_url}/attend/{event_id}"

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=8,
            border=3,
        )
        qr.add_data(attend_url)
        qr.make(fit=True)

        img = qr.make_image(fill_color="#00e5a0", back_color="#0a1520")
        stream = io.BytesIO()
        img.save(stream, format="PNG")
        stream.seek(0)

        b64 = base64.b64encode(stream.read()).decode()
        return {
            "qr_base64": f"data:image/png;base64,{b64}",
            "attend_url": attend_url,
            "event_id": event_id,
        }

    @staticmethod
    def register_attendance(event_id: int):
        """Фіксує відвідування — збільшує attendance на 1."""
        ev = Event.query.get(event_id)
        if not ev:
            return None, "Захід не знайдено"
        if ev.status not in ("Заплановано", "У процесі"):
            return None, "Реєстрація для цього заходу закрита"
        ev.attendance = (ev.attendance or 0) + 1
        db.session.commit()
        return ev.to_dict(), None


class ReportService:
    @staticmethod
    def generate_docx():
        doc = Document()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(REPORT_STRINGS["target_head"])
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(REPORT_STRINGS["doc_title"]).bold = True
        doc.add_paragraph(REPORT_STRINGS["intro_text"])
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text=REPORT_STRINGS["th_date"]; hdr[1].text=REPORT_STRINGS["th_title"]
        hdr[2].text=REPORT_STRINGS["th_loc"];  hdr[3].text=REPORT_STRINGS["th_resp"]
        for ev in Event.query.order_by(Event.date).all():
            row = table.add_row().cells
            row[0].text=str(ev.date); row[1].text=str(ev.title)
            row[2].text=str(ev.location or ""); row[3].text=str(ev.responsible or "")
        doc.add_paragraph(REPORT_STRINGS["outro_text"])
        stream = io.BytesIO(); doc.save(stream); stream.seek(0)
        return stream

    @staticmethod
    def generate_pdf():
        stream = io.BytesIO()
        doc = SimpleDocTemplate(stream, pagesize=A4,
                                rightMargin=2*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph(REPORT_STRINGS["target_head"].replace("\n","<br/>"),
                                ParagraphStyle("h", parent=styles["Normal"], fontSize=10, alignment=2, spaceAfter=12)))
        story.append(Paragraph(REPORT_STRINGS["doc_title"],
                                ParagraphStyle("t", parent=styles["Heading1"], fontSize=14, alignment=1, spaceAfter=12)))
        story.append(Paragraph(REPORT_STRINGS["intro_text"],
                                ParagraphStyle("n", parent=styles["Normal"], fontSize=10, spaceAfter=8)))
        story.append(Spacer(1, 0.3*cm))
        data = [[REPORT_STRINGS["th_date"], REPORT_STRINGS["th_title"],
                 REPORT_STRINGS["th_loc"],  REPORT_STRINGS["th_resp"], "Статус"]]
        for ev in Event.query.order_by(Event.date).all():
            data.append([ev.date, ev.title, ev.location or "", ev.responsible or "", ev.status])
        t = Table(data, colWidths=[2.5*cm, 7*cm, 4*cm, 4*cm, 2.5*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1a2a3a")),
            ("TEXTCOLOR",(0,0),(-1,0),colors.white),
            ("FONTSIZE",(0,0),(-1,0),9), ("FONTSIZE",(0,1),(-1,-1),8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#f0f4f8")]),
            ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#cccccc")),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",(0,0),(-1,-1),4), ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(REPORT_STRINGS["outro_text"],
                                ParagraphStyle("n2", parent=styles["Normal"], fontSize=10)))
        doc.build(story)
        stream.seek(0)
        return stream
